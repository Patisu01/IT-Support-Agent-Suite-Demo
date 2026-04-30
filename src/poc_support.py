from __future__ import annotations

import csv
import json
import os
import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

try:
    import openpyxl
    from docx import Document
    from langchain_community.vectorstores import FAISS
    from langchain_core.documents import Document as LangChainDocument
    from langchain_core.embeddings import Embeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from sentence_transformers import SentenceTransformer
except ModuleNotFoundError as exc:
    missing_package = exc.name or "a required package"
    raise ModuleNotFoundError(
        f"Missing dependency: {missing_package}. Install this project's Python packages with "
        "'python -m pip install -r requirements.txt' from the repository root, or run the script "
        "with the bundled runtime documented in README.md."
    ) from exc


def load_config(config_path: Path) -> dict:
    return json.loads(config_path.read_text(encoding="utf-8"))


def normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def tokenize(value: str) -> List[str]:
    return [token for token in normalize_text(value).split() if token]


@dataclass
class Ticket:
    ticket_id: str
    store_code: str
    location_label: str
    issue: str
    priority: str
    description: str


@dataclass
class RetrievalResult:
    score: float
    chunk: str


class SentenceTransformerEmbeddings(Embeddings):
    """Offline-capable wrapper around a stronger sentence-transformer model."""

    def __init__(self, model_name: str, cache_dir: Path):
        cache_dir.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault("HF_HOME", str(cache_dir))
        os.environ.setdefault("SENTENCE_TRANSFORMERS_HOME", str(cache_dir))
        os.environ.setdefault("HF_HUB_OFFLINE", "1")
        os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
        self.model = SentenceTransformer(
            model_name,
            cache_folder=str(cache_dir),
            local_files_only=True,
        )

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        vectors = self.model.encode(texts, normalize_embeddings=True)
        return vectors.tolist()

    def embed_query(self, text: str) -> List[float]:
        vector = self.model.encode(text, normalize_embeddings=True)
        return vector.tolist()


class AssetRepository:
    def __init__(self, asset_csv: Path):
        self.rows = self._load_rows(asset_csv)

    @staticmethod
    def _load_rows(asset_csv: Path) -> List[dict]:
        with asset_csv.open("r", encoding="utf-8-sig", newline="") as handle:
            return list(csv.DictReader(handle))

    def get_assets_for_store(self, store_id: str) -> List[dict]:
        return [row for row in self.rows if row["Store_ID"] == store_id]

    def get_pos_assets_for_store(self, store_id: str) -> List[dict]:
        return [row for row in self.get_assets_for_store(store_id) if row["Asset_Type"] == "POS"]


class ITSMRepository:
    def __init__(self, itsm_xlsx: Path):
        self.rows = self._load_rows(itsm_xlsx)

    @staticmethod
    def _load_rows(itsm_xlsx: Path) -> List[dict]:
        workbook = openpyxl.load_workbook(itsm_xlsx, read_only=True, data_only=True)
        worksheet = workbook[workbook.sheetnames[0]]
        iterator = worksheet.iter_rows(values_only=True)
        headers = next(iterator)
        rows = []
        for row in iterator:
            rows.append({headers[i]: row[i] for i in range(len(headers))})
        return rows

    def history_for_store(self, store_code: str) -> List[dict]:
        return [row for row in self.rows if store_code in str(row.get("Location", ""))]

    def recurring_issue_summary(self, store_code: str) -> str:
        history = self.history_for_store(store_code)
        if not history:
            return "No historical incidents found for this store."
        categories = Counter(str(row.get("Issue_Category", "")) for row in history)
        top_issue, count = categories.most_common(1)[0]
        return f"Most frequent historical issue at {store_code}: {top_issue} ({count} incidents)."


class SOPRetriever:
    def __init__(self, sop_docx: Path):
        self.mode = "langchain_faiss_sentence_transformer"
        self.documents = self._load_documents(sop_docx)
        self.vectorstore = self._build_vectorstore(self.documents, sop_docx)

    @staticmethod
    def _load_documents(sop_docx: Path) -> List[LangChainDocument]:
        document = Document(sop_docx)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=250,
            chunk_overlap=40,
            separators=["\n\n", "\n", ". ", "; ", " "],
        )

        chunks = text_splitter.split_text("\n".join(paragraphs))
        return [
            LangChainDocument(
                page_content=chunk,
                metadata={"source": str(sop_docx), "chunk_id": index},
            )
            for index, chunk in enumerate(chunks)
        ]

    @staticmethod
    def _build_vectorstore(documents: List[LangChainDocument], sop_docx: Path) -> FAISS:
        cache_dir = Path(__file__).resolve().parent.parent / "artifacts" / "model_cache"
        embeddings = SentenceTransformerEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            cache_dir=cache_dir,
        )
        return FAISS.from_documents(documents, embeddings)

    def retrieve(self, query: str, limit: int = 4) -> List[RetrievalResult]:
        matches = self.vectorstore.similarity_search_with_score(query, k=limit)
        return [
            RetrievalResult(
                score=round(float(score), 4),
                chunk=document.page_content,
            )
            for document, score in matches
        ]


class MockPlatformTools:
    def __init__(self, heartbeat_defaults: Dict[str, str]):
        self.heartbeat_defaults = heartbeat_defaults
        self.comments: List[dict] = []
        self.actions: List[dict] = []

    def check_pos_heartbeat(self, store_id: str) -> str:
        status = self.heartbeat_defaults.get(store_id, "Unknown")
        self.actions.append({"tool": "check_pos_heartbeat", "store_id": store_id, "result": status})
        return status

    def trigger_port_reset(self, asset_id: str) -> str:
        result = f"Remote port reset simulated successfully for {asset_id}."
        self.actions.append({"tool": "trigger_port_reset", "asset_id": asset_id, "result": result})
        return result

    def update_incident_comment(self, ticket_id: str, comment: str) -> str:
        self.comments.append({"ticket_id": ticket_id, "comment": comment})
        return f"Incident comment recorded for {ticket_id}."


def infer_family(model: str, family_map: Dict[str, str]) -> str:
    return family_map.get(model, "Unmapped POS family")


def choose_primary_pos_asset(pos_assets: List[dict]) -> dict | None:
    if not pos_assets:
        return None

    def rank(asset: dict) -> tuple:
        status = asset.get("Status", "")
        healthy = status not in {"Decommissioned", "In Repair"}
        return (healthy, asset.get("Model", ""))

    return sorted(pos_assets, key=rank, reverse=True)[0]


def render_history(history: List[dict], limit: int = 3) -> List[str]:
    lines = []
    for item in history[:limit]:
        ticket_id = item.get("Ticket_ID", "Unknown")
        issue = item.get("Issue_Category", "Unknown issue")
        resolution = item.get("Resolution", "No resolution logged")
        lines.append(f"{ticket_id}: {issue} -> {resolution}")
    return lines
